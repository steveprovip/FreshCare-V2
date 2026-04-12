# -*- coding: utf-8 -*-
import docx
from docx.oxml.ns import qn

doc = docx.Document(r'BaoCaoCongNghePhanMem_Nhom8_KTKN1_N08_V9 (1).docx')

# Check for SDT (structured document tags) used by TOC
body = doc.element.body
sdt_elements = body.findall('.//' + qn('w:sdt'))
print('SDT elements:', len(sdt_elements))
for sdt in sdt_elements:
    alias = sdt.find('.//' + qn('w:alias'))
    if alias is not None:
        print('SDT alias:', alias.get(qn('w:val')))

# Check TOC area
print()
for i in range(27, 35):
    p = doc.paragraphs[i]
    xml_snip = p._element.xml[:500]
    text = p.text[:100] if p.text else "[EMPTY]"
    print(f"Para {i}: style={p.style.name} text={text}")
    if 'fld' in xml_snip.lower() or 'toc' in xml_snip.lower():
        print("  >>> Contains field/TOC XML")
        print("  XML snippet:", xml_snip[:300])

# Check all heading paragraphs to build TOC structure
print("\n=== HEADING STRUCTURE ===")
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading'):
        print(f"Para {i}: [{p.style.name}] {p.text[:120]}")

# Count images/inline shapes
print(f"\nInline shapes: {len(doc.inline_shapes)}")
